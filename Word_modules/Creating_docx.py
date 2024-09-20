import json
import os
import re
from bs4 import BeautifulSoup
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.oxml import qn
from docx.oxml import OxmlElement
from tqdm import tqdm
from pdf2docx import Converter
from Word_modules.word_data import  metadata
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from openai import OpenAI
from collections import defaultdict
# Uncomment if on Windows and docx2pdf is installed
from docx2pdf import convert
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_style(style, font_name=None, font_size=None, bold=None, alignment=None, line_spacing=None, space_after=None,
              first_line_indent=None):
    """Helper function to set common style attributes."""
    font = style.font
    paragraph_format = style.paragraph_format

    if font_name:
        font.name = font_name
    if font_size:
        font.size = font_size
    if bold is not None:
        font.bold = bold
    if alignment:
        paragraph_format.alignment = alignment
    if line_spacing:
        paragraph_format.line_spacing = line_spacing
    if space_after:
        paragraph_format.space_after = space_after
    if first_line_indent:
        paragraph_format.first_line_indent = first_line_indent


def add_multilevel_list(doc):
    """Create multilevel numbering scheme in the document."""
    numbering = doc._part.numbering_part.element  # Access the numbering part of the document
    abstractNum = OxmlElement('w:abstractNum')
    abstractNum.set(qn('w:abstractNumId'), '0')

    levels = [
        ('w:lvl', 'w:start', 'w:numFmt', 'decimal', 'w:lvlText', '%1', 'w:lvlJc', 'left', 0),
        ('w:lvl', 'w:start', 'w:numFmt', 'decimal', 'w:lvlText', '%1.%2', 'w:lvlJc', 'left', 1),
        ('w:lvl', 'w:start', 'w:numFmt', 'decimal', 'w:lvlText', '%1.%2.%3', 'w:lvlJc', 'left', 2),
        ('w:lvl', 'w:start', 'w:numFmt', 'decimal', 'w:lvlText', '%1.%2.%3.%4', 'w:lvlJc', 'left', 3),
        ('w:lvl', 'w:start', 'w:numFmt', 'decimal', 'w:lvlText', '%1.%2.%3.%4.%5', 'w:lvlJc', 'left', 4),
    ]

    for lvl, start, numFmt, fmt, lvlText, text, lvlJc, jc, ilvl in levels:
        level = OxmlElement(lvl)
        level.set(qn('w:ilvl'), str(ilvl))

        start_val = OxmlElement(start)
        start_val.set(qn('w:val'), '1')
        level.append(start_val)

        numFmt_val = OxmlElement(numFmt)
        numFmt_val.set(qn('w:val'), fmt)
        level.append(numFmt_val)

        lvlText_val = OxmlElement(lvlText)
        lvlText_val.set(qn('w:val'), text)
        level.append(lvlText_val)

        lvlJc_val = OxmlElement(lvlJc)
        lvlJc_val.set(qn('w:val'), jc)
        level.append(lvlJc_val)

        abstractNum.append(level)

    numbering.append(abstractNum)

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), '1')
    abstractNumId = OxmlElement('w:abstractNumId')
    abstractNumId.set(qn('w:val'), '0')
    num.append(abstractNumId)
    numbering.append(num)


def assign_numbering_to_style(doc, style_name, level):
    """Assign multilevel numbering to a style."""
    style = doc.styles[style_name]
    pPr = style.element.get_or_add_pPr()

    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')  # The numId corresponding to the list definition
    numPr.append(ilvl)
    numPr.append(numId)

    pPr.append(numPr)


def create_numbered_heading_style(doc):
    """Create numbered heading styles and apply multilevel list."""
    # Add multilevel list to the document for numbering
    add_multilevel_list(doc)

    # Define numbered heading styles (Heading 1, Heading 2, ..., Heading 5)
    for level in range(5):  # Now supporting up to Heading 5
        style_name = f'Heading {level + 1}'
        if style_name not in doc.styles:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[style_name]

        set_style(style, font_name='Times New Roman', font_size=Pt(14 - level), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
        # Assign numbering to the heading style
        assign_numbering_to_style(doc, style_name, level)


def customize_doc_style(doc, header_title, block_quote_indent=Inches(0.5), left_margin=Inches(1.5)):
    # Set margins (1 inch all around, 1.5 inches on the left if needed for binding)
    # Create a new paragraph style called 'Custom TOC Heading'
    custom_heading_style = doc.styles.add_style('topic_sentence', WD_STYLE_TYPE.PARAGRAPH)

    # Customize the appearance of 'Custom TOC Heading'
    set_style(custom_heading_style, font_name='Times New Roman', font_size=Pt(14), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Set the custom style's outline level (e.g., equivalent to Heading 1 in TOC)
    custom_heading_style.paragraph_format.outline_level = 1  # Equivalent to 'Heading 1' in TOC
def add_page_number(paragraph):
    """Add page number to the specified paragraph."""
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # Create a new field character
    fldChar.set(qn('w:fldCharType'), 'begin')  # Field type: begin
    instrText = OxmlElement('w:instrText')  # Create instruction element
    instrText.text = 'PAGE'  # This is the field code for page number
    fldChar_end = OxmlElement('w:fldChar')  # Create end of field character
    fldChar_end.set(qn('w:fldCharType'), 'end')  # Field type: end
    run._r.append(fldChar)  # Add begin
    run._r.append(instrText)  # Add instruction
    run._r.append(fldChar_end)  # Add end


def customize_doc_style(doc, header_title, block_quote_indent=Inches(0.5), left_margin=Inches(1.5)):
    # Set margins (1 inch all around, 1.5 inches on the left if needed for binding)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = left_margin
        section.right_margin = Inches(1)

    # Customize 'Normal' style for body text in academic papers (1.5 line spacing)
    normal_style = doc.styles['Normal']
    set_style(normal_style, font_name='Times New Roman', font_size=Pt(12),
              line_spacing=1.5, space_after=Pt(0), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
              first_line_indent=Inches(0.5))  # 1.5 line spacing with first-line indent of 0.5 inches

    # Add automatic page numbering, skipping the first page (title page)
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(paragraph)

    # Customize Header style with the title, apply only after the first page
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.text = header_title
    set_style(paragraph.style, font_name='Times New Roman', font_size=Pt(12),
              alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Add Block Quote style (indented, single-spaced)
    if 'Block Quote' in doc.styles:
        block_quote_style = doc.styles['Block Quote']
    else:
        block_quote_style = doc.styles.add_style('Block Quote', WD_STYLE_TYPE.PARAGRAPH)

    set_style(block_quote_style, font_name='Times New Roman', font_size=Pt(12),
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.0, first_line_indent=None)
    block_quote_style.paragraph_format.left_indent = block_quote_indent

    # Add numbered heading styles globally
    create_numbered_heading_style(doc)

    # Add title page formatting (optional)
    title_page_section = doc.sections[0]
    title_page_section.start_type = WD_SECTION.NEW_PAGE  # Ensure it's a new section

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run(header_title)
    run.bold = True
    run.font.size = Pt(16)

    # Ensure the document starts on a new page after the title page
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Ensure use of section breaks when needed
    for paragraph in doc.paragraphs:
        # Manually add a section break at the end of each major section if required
        if paragraph.text.strip().lower() in ['introduction', 'conclusion', 'references']:
            doc.add_section(WD_SECTION.NEW_PAGE)  # Add a new section starting on a new page


    # Customize the References section style with hanging indent
    reference_style = doc.styles.add_style('References', WD_STYLE_TYPE.PARAGRAPH)
    set_style(reference_style, font_name='Times New Roman', font_size=Pt(12),
              line_spacing=1.5, space_after=Pt(12), alignment=WD_ALIGN_PARAGRAPH.LEFT)
    reference_style.paragraph_format.first_line_indent = Inches(-0.5)  # Hanging indent of 0.5 inches


class Docx_creation():
    def __init__(self):
        pass

    def save_docx(self, directory_path,output_path):
        # Append ".json" to create the complete file path
        with open(output_path, "r",encoding="utf-8") as dt:
                data = json.load(dt)
        output_path = os.path.split(output_path)[-1]
        file_path = f"{directory_path}.docx"
        doc = Document()
        self.set_document_metadata(doc)
        self.add_cover_page(doc)
        # doc.add_heading("Introduction", level=1)
        # intro = doc.add_paragraph(introduction_content["Introduction"])
        # intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        # doc.add_paragraph()  # Optionally add extra spacing between sections
        # doc.add_page_break()

        self.customize_doc_style(doc)  # Apply the custom styles to the document

        chapter_name = list(data.keys())[0]
        chapter_title = data[chapter_name]["title"]

        for chapter_section in data[chapter_name]["Sections"]:


            section_title = chapter_section["title"]
            section_intro =  chapter_section["introduction"]
            section_keywords = chapter_section["keywords"]
            subsections = chapter_section["Subsections"]

            # Use 'Heading 1' for main headings
            doc.add_heading(chapter_title.strip().title(), level=1).style = 'Heading 1'
            doc.add_heading(section_title.strip().title(), level=2).style = 'Heading 2'

            intro_text = doc.add_paragraph(section_intro)
            intro_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            with tqdm(chapter_section["Subsections"], desc=f'Processing {section_title}') as pbar:
                for n, subsection in enumerate(pbar):
                    subsection_title = subsection["title"].strip().title()
                    sub_text = subsection.get("bodyText").strip() if subsection.get("bodyText") else ""
                    doc.add_heading(subsection_title, level=3).style = 'Heading 3'

                    subsection_ref = "\n".join(subsection["references"])
                    print("subsection ref:", subsection_ref)



                    body_text = doc.add_paragraph(sub_text)
                    body_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    doc.add_heading("Reference_processing", level=2).style = 'Heading 3'
                    ref_text = doc.add_paragraph(subsection_ref)
                    ref_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY





            doc.add_page_break()

        doc.save(file_path)
        print("the docx file has been saved at {}\n".format(file_path))
        convert(file_path, file_path.replace(".docx", ".pdf"))
        print("the pdf file has been saved at {}\n".format(file_path.replace(".docx", ".pdf")))


    def create_scientific_template(self, doc):
        # Set up document styles (fonts, sizes, line spacing)
        styles = doc.styles
        normal = styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(12)

        for i in range(1, 4):  # Create Heading styles
            heading_style = styles.add_style(f'Heading {i}', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = 'Times New Roman'
            heading_style.font.size = Pt(12 + 2 * (3 - i))
            heading_style.font.bold = True
            heading_style.paragraph_format.space_after = Pt(6)

        # Create a title page
        doc.add_heading('Research Title', 0)
        doc.add_paragraph('Author Name\nInstitution\nDate').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # Add an abstract
        doc.add_heading('Abstract', level=1)
        doc.add_paragraph('This is a placeholder for the abstract.')

        # Add keywords
        doc.add_paragraph('Keywords: keyword1, keyword2, keyword3')

        # Introduction
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph('This is a placeholder for the introduction.')

        # Methodology
        doc.add_heading('Methodology', level=1)
        doc.add_paragraph('This is a placeholder for the methodology.')

        # Results
        doc.add_heading('Results', level=1)
        doc.add_paragraph('This is a placeholder for the results.')

        # Discussion
        doc.add_heading('Discussion', level=1)
        doc.add_paragraph('This is a placeholder for the discussion.')

        # Reference_processing
        doc.add_heading('Reference_processing', level=1)
        doc.add_paragraph('This is a placeholder for the references.')

        # Appendices
        doc.add_heading('Appendices', level=1)
        doc.add_paragraph('This is a placeholder for the appendices.')

    def set_document_metadata(self, doc):
        """
        Set the document's core properties (metadata).
        :param doc: The Document object.
        :param metadata: A dictionary containing metadata fields.
        """
        core_properties = doc.core_properties
        core_properties.author = metadata.get('author', 'Unknown Author')
        core_properties.title = metadata.get('title', 'Untitled Document')
        core_properties.subject = metadata.get('subject', '')
        core_properties.keywords = metadata.get('keywords', '')
        core_properties.comments = metadata.get('comments', '')
        core_properties.category = metadata.get('category', 'General')
        core_properties.version = metadata.get('version', '1.0')
        core_properties.language = metadata.get('language', 'en-US')
        core_properties.created = metadata.get('created', datetime.now())
        core_properties.modified = metadata.get('modified', datetime.now())
        # Add any additional metadata as needed


    def add_cover_page(self, doc):
        """
        Add a cover page to the document.
        :param doc: The Document object.
        :param metadata: A dictionary containing metadata fields.
        """
        doc.add_paragraph(metadata['title'], style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph('\n' * 5)  # Adjust spacing as needed
        doc.add_paragraph(metadata['author'], style='Subtitle').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Ensure date is a datetime object for formatting
        date_str = metadata['created'].strftime('%B %d, %Y') if isinstance(metadata['created'], datetime) else metadata[
            'created']
        doc.add_paragraph('\n' * 13)  # Move date & city to bottom
        doc.add_paragraph(f"{date_str}, {metadata['city']}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_page_break()


        # # Add page number to the footer
        # fldSimple = OxmlElement('w:fldSimple')
        # fldSimple.set(qn('w:instr'), 'PAGE')
        # run = paragraph.add_run()
        # run._r.append(fldSimple)


    # docx = Docx_creation()
    # docx.save_docx(directory_path="Cyber Evidence", output_path="books_data/Cyber Evidence/Cyber evidence.json")

    def create_word_document_from_html(self,html_content, output_path):
        # Initialize a new Word document
        doc = Document()

        # Parse HTML content
        soup = BeautifulSoup(html_content, 'html.parser')

        # Extract headings and paragraphs
        headings = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
        paragraphs = soup.find_all('p')

        # Add headings to the document
        for heading in headings:
            text = heading.get_text()
            level = int(heading.name[1])  # Extract heading level from tag name (e.g., 'h1' -> level 1)
            doc.add_heading(text, level=level)

        # Add paragraphs to the document
        for paragraph in paragraphs:
            text = paragraph.get_text()
            doc.add_paragraph(text)

        # Save the document
        doc.save(output_path)

    #
    # output_path = "output.docx"
    # create_word_document_from_html(html_content, output_path)

    def add_texts_to_html(self,html_path, headings_texts_dict, start_from=None):
        with open(html_path, 'r') as file:
            html_content = file.readlines()

        updated_html_content = []
        processing = not start_from  # If start_from is None, start processing immediately

        for line in html_content:
            if not processing:
                if start_from in line:
                    processing = True
                else:
                    updated_html_content.append(line.strip() + '\n')
                    continue

            for heading, text in headings_texts_dict.items():
                if heading in line:
                    updated_html_content.append(line.strip() + '\n')
                    updated_html_content.append(text.strip() + '\n')
                    break
            else:
                updated_html_content.append(line.strip() + '\n')

        with open(html_path, 'w') as file:
            file.writelines(updated_html_content)

    def pdf_to_docx(self,pdf_path):
        import logging
        # Setting up logging
        logging.basicConfig(filename='pdf_conversion.log', level=logging.DEBUG)

        # Extracting the file name without extension and the directory path
        docx_path = pdf_path.rsplit('.', 1)[0] + '.docx'

        # Create a Converter object and perform conversion
        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None)
            cv.close()
            return docx_path
        except Exception as e:
            logging.error(f"Failed to convert {pdf_path}: {e}")

            return None

    # Example usage:
    # pdf_path = r"C:\Users\luano\Zotero\storage\LHZRBA2H\Johnson and Schmitt - 2021 - Responding to proxy cyber operations under international law.pdf" # Replace with your PDF file path
    # pdf_to_docx(pdf_path)
    # from docx import Document

    from docx import Document
    import os
    def extract_paragraphs_and_footnotes(self, pdf_path, quote,author,stop_words):
        docx_path = pdf_path.replace('.pdf', '.docx')
        if os.path.exists(docx_path):
            print("The DOCX file exists.")
        else:
            self.pdf_to_docx(pdf_path)  # Assume this is a method defined elsewhere in the class
            print("The DOCX file does not exist.")

        doc = Document(docx_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        merged_paragraphs = []
        footnotes = {}
        current_paragraph = ""
        current_footnote_number = 0
        current_footnote_content = ""

        # Regular expressions to detect repetitive and unwanted content
        unwanted_patterns = [
            r'^This content downloaded from \d+\.\d+\.\d+\.\d+ on \w+, \d+ \w+ \d{4} \d{2}:\d{2}:\d{2} \+\d{4} All use subject to https://about\.jstor\.org/terms$',
            r'This content downloaded from 92.40.169.176 on Fri, 10 May 2024 08:35:35 +00:00 All use subject to https:',  # Pagination pattern
            r'RESPONDING TO PROXY CYBER OPERATIONS UNDER INTERNATIONAL LAW',
            r'^\d...THE CYBER DEFENSE REVIEW',
            r'DURWARD E. JOHNSON : MICHAEL N. SCHMITT',
        ]
        len_paragraphs = len(paragraphs)
        for i, paragraph in enumerate(paragraphs):
            # Skip unwanted repetitive lines or pagination
            if any(re.match(pattern, paragraph) for pattern in unwanted_patterns):
                continue

            # Handle footnotes
            if paragraph[0].isdigit() and '.' in paragraph[:3]:
                potential_number = paragraph.split('.')[0]
                if potential_number.isdigit():
                    current_footnote_number = int(potential_number)
                    current_footnote_content = paragraph[len(potential_number) + 1:].strip()
                    if current_footnote_content.endswith('.'):
                        footnotes[current_footnote_number] = current_footnote_content
                        current_footnote_content = ""
                    continue

            # Remove single-word or very short lines
            if len(paragraph.split()) <= 2:
                continue

            # Special handling for paragraphs that may end with period followed by a footnote marker
            if i + 1 < len_paragraphs and re.search(r'\[\d+\]$', paragraph) and paragraphs[i + 1][0].isupper():
                if current_paragraph:
                    merged_paragraphs.append(current_paragraph)
                current_paragraph = paragraph
            elif current_footnote_content:
                current_footnote_content += " " + paragraph
                if paragraph.endswith('.'):
                    footnotes[current_footnote_number] = current_footnote_content
                    current_footnote_content = ""
            else:
                # Ensure paragraphs start with a capital and are properly ended
                if paragraph[0].isupper() and (paragraph.endswith('.') or re.search(r'\[\d+\]$', paragraph)):
                    if current_paragraph:
                        merged_paragraphs.append(current_paragraph)
                    current_paragraph = paragraph
                else:
                    current_paragraph += " " + paragraph

        # Append any remaining content
        if current_paragraph:
            merged_paragraphs.append(current_paragraph)
        if current_footnote_content:
            footnotes[current_footnote_number] = current_footnote_content

        # Match and interact based on a quote
        for para in merged_paragraphs:
            if quote[4:30].replace('"', '').replace("'", "") in para.replace('"', '').replace("'", ""):
                print(quote[4:30])
                print(para)
                print("\n\n\n")
                return para
        # return {'paragraphs': merged_paragraphs, 'footnotes': footnotes}