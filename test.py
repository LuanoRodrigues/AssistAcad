import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
import pygetwindow as gw
from pywinauto import Application
import time
import statistics

from pdf2docx import Converter

def pdf_to_docx(pdf_path):
    # Extracting the file name without extension and the directory path
    docx_path = pdf_path.rsplit('.', 1)[0] + '.docx'

    # Create a Converter object and perform conversion
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()

    print(f"Converted '{pdf_path}' to '{docx_path}'")

# Example usage:
# pdf_path = r"C:\Users\luano\Zotero\storage\LHZRBA2H\Johnson and Schmitt - 2021 - Responding to proxy cyber operations under international law.pdf" # Replace with your PDF file path
# pdf_to_docx(pdf_path)
# from docx import Document


from docx import Document
import os
def extract_paragraphs_and_footnotes(docx_path,quote):


    if os.path.exists(docx_path):
        print("The DOCX file exists.")
    else:
        pdf_to_docx(docx_path.replace('.docx','.pdf'))
        print("The DOCX file does not exist.")
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    merged_paragraphs = []
    footnotes = {}
    current_paragraph = ""
    current_footnote_number = 0
    current_footnote_content = ""

    for paragraph in paragraphs:
        # Detect if the paragraph is potentially a footnote by checking if it starts with a number followed by a period
        if paragraph[0].isdigit():
            potential_number = paragraph.split()[0][:-1]
            if potential_number.isdigit():  # Confirms it's a simple number followed by a period
                current_footnote_number = int(potential_number)
                current_footnote_content = paragraph[len(potential_number)+1:].strip()
                if current_footnote_content.endswith('.'):
                    footnotes[current_footnote_number] = current_footnote_content
                    current_footnote_content = ""
                continue

        # Handle continuation of footnotes
        if current_footnote_content:
            current_footnote_content += " " + paragraph
            if paragraph.endswith('.'):
                footnotes[current_footnote_number] = current_footnote_content
                current_footnote_content = ""
        else:
            # Handle paragraphs that must not start with numbers and ensure they end with a period
            if current_paragraph and (not paragraph[0].isdigit() and (paragraph[0].islower() or not current_paragraph.endswith('.'))):
                current_paragraph += " " + paragraph
            else:
                if current_paragraph:
                    merged_paragraphs.append(current_paragraph)
                current_paragraph = paragraph

    # Append any remaining content
    if current_paragraph:
        if quote in current_paragraph:

            return current_paragraph
        merged_paragraphs.append(current_paragraph)
    if current_footnote_content:  # Append any remaining footnote content
        footnotes[current_footnote_number] = current_footnote_content
    for para in merged_paragraphs:
        if quote.replace('"', '').replace("'", "") in para.replace('"', '').replace("'", ""):
            return para
    # return {'paragraphs': merged_paragraphs, 'footnotes': footnotes}


from docx import Document


def extract_non_body_text(doc_path):
    doc = Document(doc_path)
    non_body_text = []

    for paragraph in doc.paragraphs:
        if 'https://' in paragraph.text or 'www.' in paragraph.text:
            non_body_text.append(paragraph.text)
        elif 'Vol.' in paragraph.text or 'No.' in paragraph.text:
            non_body_text.append(paragraph.text)
        elif 'pp.' in paragraph.text or 'Published by:' in paragraph.text:
            non_body_text.append(paragraph.text)
        elif 'doi:' in paragraph.text or 'DOI:' in paragraph.text:
            non_body_text.append(paragraph.text)
        elif paragraph.text.isupper():
            non_body_text.append(paragraph.text)
        elif paragraph.style.name.startswith('Heading'):
            non_body_text.append(paragraph.text)
        elif any(char.isdigit() for char in paragraph.text) and len(paragraph.text) < 20:
            non_body_text.append(paragraph.text)

    return non_body_text


def clean_text(doc_path, clean_list):
    doc = Document(doc_path)
    for paragraph in doc.paragraphs:
        for item in clean_list:
            if item in paragraph.text:
                paragraph.text = paragraph.text.replace(item, '')
    return doc


def main(doc_path):
    # Extract non-body text elements
    non_body_text = extract_non_body_text(doc_path)

    # Clean the document
    cleaned_doc = clean_text(doc_path, non_body_text)

    # Save the cleaned document
    cleaned_doc.save('cleaned_document.docx')
    print(cleaned_doc)
    print('The document has been cleaned and saved as cleaned_document.docx')


# Example usage
docx_path = r"C:\Users\luano\Zotero\storage\LHZRBA2H\Johnson and Schmitt - 2021 - Responding to proxy cyber operations under international law.docx"  # Replace with your DOCX file path
# result = extract_paragraphs_and_footnotes(docx_path,"The “direction or control” standard applies when the proxy’s affiliation with")
# print(result)

main(docx_path)