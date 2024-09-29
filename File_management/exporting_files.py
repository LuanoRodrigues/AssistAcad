import json
import os.path
import zipfile
import json
from typing import List
from bs4 import BeautifulSoup, NavigableString
from typing import List, Tuple
import markdown
import markdownify
import pdfkit
from bs4 import BeautifulSoup
from docx import Document
from docx2pdf import convert
import json
from docx import Document
from scipy.signal import impulse
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE

from Pychat_module.gpt_api import call_openai_api

from Word_modules.Creating_docx import customize_doc_style
md_root=r"C:\Users\luano\OneDrive - University College London\Obsidian\cyber evidence"

def export_results(structure, filename="output", export_to=['doc','html']):
    export_to = export_to or ['doc', 'html', 'md', 'json']
    doc = Document()

    # Export as HTML
    if "html" in export_to:
        with open(f"{filename}.html", "w", encoding='utf-8') as html_file:
            overarching_theme = structure.get('overarching_theme', 'No overarching theme')
            html_file.write(f"<h1>{overarching_theme}</h1>\n")

            for item in structure.get('structure', []):
                print(item)
                if isinstance(item, str):
                    item=eval(item)
                level =int(item['heading'].replace('h', '').strip())

                if isinstance(item, str):  # Check if item is a string (JSON)
                    item = json.loads(item)  # Parse the JSON string into a dictionary

                if 'heading' in item.keys() and 'paragraph_ids' in item.keys():
                    if 'topic_sentence' in item.keys():
                        quote_tag=f'h{str(level+1)}'
                        html_file.write(f"<{item['heading']}>{item['title']}</{item['heading']}\n")
                        for i in range(len(item['paragraph_ids'])):
                            topic=item['topic_sentence'][i]
                            paragraph=item['paragraph_ids'][i]

                            html_file.write(f"<{quote_tag}>{topic}</{quote_tag}>\n")
                            html_file.write(f"<p>{paragraph}</p>\n")

                else:
                        html_file.write(f"<{item['title']}>{item['title']}</{item['title']}\n")
                        for paragraph in item['paragraphs']:
                            html_file.write(f"<p>{paragraph}</p>\n")


    # Export as DOCX
    if "doc" in export_to:

        print('yes')
        customize_doc_style(doc, 'header here')

        # Add the overarching theme as a level-1 heading
        doc.add_heading(structure.get('overarching_theme', 'No overarching theme'), level=1)

        for item in structure.get('structure', []):
            if isinstance(item, str):  # Check if item is a string (JSON)
                item = json.loads(item)  # Parse the JSON string into a dictionary

            if 'heading' in item.keys() and 'title' in item.keys() :

                # Combine the heading level and title for the section
                heading_text = f"{item['title']}"
                # Add the heading to the doc with the title as level-2 heading (or the appropriate level)
                # Adjust 'level=2' if needed based on your structure
                level =int(item['heading'].replace('h', '').strip())
                doc.add_heading(heading_text, level=level)
                # Add the paragraphs
                if 'topic_sentence' in item:
                    for i in range(len(item['paragraph_ids'])):
                        topic = item['topic_sentence'][i]
                        paragraph = item['paragraph_ids'][i]

                        doc.add_heading(text=topic, level=level+1)
                        doc.add_paragraph(paragraph,style='Quote')
                else:
                    for paragraph in item['paragraphs']:
                        doc.add_paragraph(paragraph)

        # Save the document
        doc.save(f"{filename}.docx")

    # Export as Markdown (MD)
    if "md" in export_to:
        with open(f"{filename}.md", "w", encoding='utf-8') as md_file:
            overarching_theme = structure.get('overarching_theme', 'No overarching theme')
            md_file.write(f"# {overarching_theme}\n\n")

            for item in structure.get('structure', []):
                if isinstance(item, str):  # Check if item is a string (JSON)
                    item = json.loads(item)  # Parse the JSON string into a dictionary

                if 'heading' in item and 'paragraphs' in item:
                    md_file.write(f"## {item['heading']}\n\n")
                    for paragraph in item['paragraphs']:
                        md_file.write(f"{paragraph}\n\n")

    # Export as JSON
    if "json" in export_to:
        with open(f"{filename}.json", "w", encoding='utf-8') as json_file:
            json.dump(structure, json_file, ensure_ascii=False, indent=4)

# Mapping heading levels to DOCX heading styles
level_map = {
    'H1': 1,
    'H2': 2,
    'H3': 3,
    'H4': 4,
    'H5': 5,
    'H6': 6
}

def add_paragraphs(doc, item):
    """Helper function to add paragraphs and blockquotes."""
    paragraph_titles = item.get('paragraph_title', [])
    paragraph_texts = item.get('paragraph_text', [])

    metadata=item.get('metadata', [])
    biblio=[]
    # Add paragraphs if present
    if paragraph_titles:
        for i, title in enumerate(paragraph_titles):
            doc.add_paragraph(title, style='IntenseQuote')  # Add title with style
            if paragraph_texts and i < len(paragraph_texts):
                doc.add_paragraph(paragraph_texts[i], style='Quote')  # Add text with style
            if metadata:
                # Use list comprehension to format each entry in metadata
                references = [
                    f"{entry.get('authors', 'Unknown Author')}. "
                    f"({entry.get('date', 'Unknown Year')}). "
                    f"{entry.get('title', 'Unknown Title')}."
                    f"{f' {entry.get("journal")}' if entry.get('journal') else ''}"
                    for entry in metadata
                ]
                print('references\n', references)
                biblio.extend(references)
    return  biblio

def process_outline(doc, item,single_theme=False):
    """Recursively process the outline and subheadings/subsections."""
    title = item.get('title', '')
    level_str = item.get('level', '')
    level = level_map.get(level_str.upper(), None)
    bibliography =[]

    # Add heading based on the level
    if title and level is not None:
        if single_theme and level_str.upper() == 'H1':
            customize_doc_style(doc, title)
        if single_theme and level_str.upper() != 'H1':

            level = level_map.get(level_str.upper(), None)-1
            doc.add_heading(title, level=level)
        else:
            doc.add_heading(title, level=level)

    # Add paragraphs (if any)
    if 'paragraph_title' in item or 'paragraph_text' in item:
        bibliography.extend(add_paragraphs(doc, item))


    # Check for subheadings and process recursively
    subheadings = item.get('subheadings', [])
    if isinstance(subheadings, list):
        for subheading in subheadings:
            bibliography.extend(process_outline(doc, subheading,single_theme=single_theme)) # Recursively process subheadings

    return bibliography

def create_docx(filename, themes):
    """Main function to create DOCX from themes and outlines."""
    bibliographic= []
    doc = Document()
    if len(themes)==1:
        customize_doc_style(doc, themes[0]['theme'])
        for outline_item in themes[0].get('outline', []):
           bibliographic.extend(process_outline(doc, outline_item,single_theme=False))  # Process each item in the outline

    if len(themes)>1:

        customize_doc_style(doc,'Thematic review')
        for theme in themes:
            for outline_item in theme.get('outline', []):
                bibliographic.append(process_outline(doc, outline_item))  # Process each item in the outline
        # Add a "References" heading
    references_heading = doc.add_paragraph("References", style='Heading 1')
    references_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add each reference with numbering
    bibliographic= sorted(set(bibliographic), key=lambda ref: ref.split(",")[0])
    for i, reference in enumerate(bibliographic, 1):
        numbered_reference = f"{i}. {reference}"  # Add numbering to each reference
        paragraph = doc.add_paragraph(numbered_reference, style='References')
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)  # Hanging indent
        paragraph.paragraph_format.left_indent = Inches(
            0.5)  # Indent the entire paragraph to preserve the hanging indent
    doc.save(filename)


    convert(filename)


def create_md(filename, themes):
    """Main function to create Markdown from themes and outlines."""
    filename= os.path.join(md_root,filename)
    md_content = ""
    bibliography = []

    for theme in themes:
        md_content += f"# {theme['theme']}\n\n"  # Add the main theme as H1
        for outline in theme.get('outline', []):
            md_content += f"## {outline.get('title', '')}\n"  # Add the outline's main title as H2

            # Add paragraphs and subheadings recursively
            outline_content, outline_biblio = add_md_paragraphs(outline)
            md_content += outline_content
            bibliography.extend(outline_biblio)

            if 'subheadings' in outline:
                subheadings_content, sub_biblio = add_md_subheadings(outline['subheadings'], level=3)
                md_content += subheadings_content
                bibliography.extend(sub_biblio)

    # Add a "References" section with numbering
    if bibliography:
        md_content += "\n# References\n\n"
        bibliography = sorted(set(bibliography))  # Remove duplicates and sort
        for i, ref in enumerate(bibliography, 1):
            md_content += f"{i}. {ref}\n\n"

    # Save the generated Markdown content to a file
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(md_content)


def add_md_paragraphs(item):
    """Helper function to add paragraphs and blockquotes to markdown content."""
    paragraph_titles = item.get('paragraph_title', [])
    paragraph_texts = item.get('paragraph_blockquote', [])
    metadata = item.get('metadata', [])
    biblio = []

    md_content = ""

    # Add paragraph titles and texts if present
    if paragraph_titles:
        for i, title in enumerate(paragraph_titles):
            md_content += f"**{title}**\n\n"  # Bold title in Markdown format
            if paragraph_texts and i < len(paragraph_texts):
                md_content += f"{paragraph_texts[i]}\n\n"  # Add corresponding text

        if metadata:
            # Use list comprehension to format metadata references in markdown
            references = [
                f"{entry.get('authors', 'Unknown Author')}. "
                f"({entry.get('date', 'Unknown Year')}). "
                f"{entry.get('title', 'Unknown Title')}."
                f"{f' *{entry.get("journal", "")}*' if entry.get('journal') else ''}"
                for entry in metadata
            ]
            biblio.extend(references)

    return md_content, biblio


def add_md_subheadings(subheadings, level=3):
    """Recursively adds subheadings and paragraphs to markdown content."""
    md_content = ""
    bibliography = []

    for sub in subheadings:
        md_heading = "#" * level
        title = sub.get('title', sub.get('heading', ''))
        md_content += f"{md_heading} {title}\n\n"  # Add heading with Markdown syntax

        # Add paragraphs under the subheading
        paragraphs_content, biblio = add_md_paragraphs(sub)
        md_content += paragraphs_content
        bibliography.extend(biblio)

        # Recursively handle sub-subheadings
        if 'subheadings' in sub:
            sub_content, sub_biblio = add_md_subheadings(sub['subheadings'], level + 1)
            md_content += sub_content
            bibliography.extend(sub_biblio)

    return md_content, bibliography



def create_html(filename, themes):
    """Main function to create HTML from themes and outlines."""
    md_content = ""

    for theme in themes:
        md_content += f"# {theme['theme']}\n\n"  # Add theme as H1
        for outline in theme.get('outline', []):
            md_content += f"## {outline.get('title', '')}\n\n"  # Add main outline title as H2

            # Add paragraphs and subheadings recursively
            outline_content, _ = add_md_paragraphs(outline)
            md_content += outline_content

            if 'subheadings' in outline:
                sub_content, _ = add_md_subheadings(outline['subheadings'], level=3)
                md_content += sub_content

    # Convert the Markdown content to HTML
    html_content = markdown.markdown(md_content, extensions=['fenced_code', 'tables'])

    # Save the HTML content to a file
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(html_content)


def create_json(filename, themes):
    with open(filename, 'w') as f:
        json.dump(themes, f, indent=4)

def export_data(filename, themes, export_to=['docx','md','html','json']):
    if 'docx' in export_to:
        create_docx(f"{filename}.docx", themes)
    if 'md' in export_to:
        create_md(f"{filename}.md", themes)
    if 'html' in export_to:
        create_html(f"{filename}.html", themes)

    if 'json' in export_to:
        create_json(f"{filename}.json", themes)


import re
import urllib.parse


def convert_zotero_to_md(html_paragraph):
    # Decode the HTML-encoded string
    decoded_paragraph = urllib.parse.unquote(html_paragraph)

    # Extract the annotation data
    annotation_regex = re.search(r'data-annotation="([^"]+)"', decoded_paragraph)
    if annotation_regex:
        annotation_data = urllib.parse.unquote(annotation_regex.group(1))
        annotation_uri_match = re.search(r'"attachmentURI":"([^"]+)"', annotation_data)
        page_match = re.search(r'"pageLabel":"([^"]+)"', annotation_data)

        if annotation_uri_match and page_match:
            annotation_uri = annotation_uri_match.group(1)
            page_number = page_match.group(1)
            annotation_link = f"[Go to annotation](zotero://open-pdf/library/items/{annotation_uri.split('/')[-1]}?page={page_number}&annotation=undefined)"
        else:
            annotation_link = ""
    else:
        annotation_link = ""

    # Extract the highlighted text
    highlight_regex = re.search(r'>([^<]+)</span>', decoded_paragraph)
    highlighted_text = highlight_regex.group(1) if highlight_regex else ""

    # Extract the citation information (including author and year)
    citation_regex = re.search(r'data-citation="([^"]+)"', decoded_paragraph)
    if citation_regex:
        citation_data = urllib.parse.unquote(citation_regex.group(1))
        citation_uri_match = re.search(r'"uris":\["([^"]+)"\]', citation_data)
        locator_match = re.search(r'"locator":"([^"]+)"', citation_data)

        # Now extract the author and year dynamically
        author_regex = re.search(r'<span class="citation-item">([^,]+)', decoded_paragraph)
        year_regex = re.search(r', (\d{4})', decoded_paragraph)

        if citation_uri_match and locator_match and author_regex and year_regex:
            citation_uri = citation_uri_match.group(1)
            page_locator = locator_match.group(1)
            author = author_regex.group(1)
            year = year_regex.group(1)
            citation_link = f"[{author}, {year}, p. {page_locator}](zotero://select/library/items/{citation_uri.split('/')[-1]})"
        else:
            citation_link = ""
    else:
        citation_link = ""

    # Combine all the extracted parts into the final markdown format
    md_output = f"{annotation_link} {highlighted_text} ({citation_link})"

    return md_output

#aaaa

# Define the heading level map
level_map = {
    'h1': 1,
    'h2': 2,
    'h3': 3,
    'h4': 4,
    'h5': 5,
    'h6': 6
}

def html_to_text_with_level(html_content: str) -> List[Tuple[str, int]]:
    """
    Convert HTML content to plain text and detect header level using BeautifulSoup.
    Returns a list of tuples [(text, level)] where level is from 1 to 6 for H1-H6 tags, 0 for paragraphs.
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    elements = []
    level_map = {'h1': 1, 'h2': 2, 'h3': 3, 'h4': 4, 'h5': 5, 'h6': 6}

    # Define the tags to process
    target_tags = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']

    # Iterate over each target tag in the content
    for tag in soup.find_all(target_tags):
        tag_name = tag.name.lower()
        texts = []

        for content in tag.descendants:
            if isinstance(content, NavigableString):
                # Append regular text
                texts.append(content.strip())
            # elif content.name == 'ref':
            #     # Append reference text inline, wrapped in parentheses
            #     ref_text = content.get_text(strip=True)
            #     # Optionally, format the reference as needed (e.g., add parentheses or superscript)
            #     texts.append(f"({ref_text})")

        # Join all parts with a space to ensure proper spacing
        text = ' '.join(filter(None, texts)).strip()
        level = level_map.get(tag_name, 0)
        elements.append((text, level))

    return elements
def export_to_word_written(data: List[dict], filename: str):
    """
    Export the data to a Word file with correct heading levels and paragraphs, preserving inline references.
    """
    data=clean_themes(data)
    for theme in data:
        outline =theme['outline']

        doc = Document()
        customize_doc_style(doc,header_title=theme['theme'])
        for entry in outline:
            section = entry.get('section', {})

            # Writing Introduction
            if 'introduction' in section:
                intro_elements = html_to_text_with_level(section['introduction'])
                for text, level in intro_elements:
                    if level > 0:
                        doc.add_heading(text, level=level)
                    else:
                        doc.add_paragraph(text, style='Normal')

            # Writing Body Text
            if 'body_text' in section:
                for body in section['body_text']:
                    body_elements = html_to_text_with_level(body)
                    for element_text, element_level in body_elements:

                        if element_level > 0:

                            doc.add_heading(element_text, level=element_level)
                        else:
                            doc.add_paragraph(element_text, style='Normal')

            # Writing Conclusion
            if 'conclusion' in section:
                conclusion_elements = html_to_text_with_level(section['conclusion'])
                for text, level in conclusion_elements:
                    if level > 0:
                        doc.add_heading(text, level=level)
                    else:
                        doc.add_paragraph(text, style='Normal')

        # Save the document to the specified filename
        doc.save(f"{filename}.docx")

def export_to_markdown(data: List[dict]):
    """
    Export the data to Markdown files by converting HTML to Markdown.
    """
    import os
    from bs4 import BeautifulSoup
    import re

    def process_ref_tags(paragraph):
        """
        Replace <ref> tags in the paragraph with the desired format.
        """
        def replace_ref(match):
            ref_tag = match.group(0)
            soup = BeautifulSoup(ref_tag, 'html.parser')
            ref = soup.find('ref')
            if ref:
                ref_id = ref.get('id', '')
                ref_text = ref.get_text(strip=True)
                # Extract author, year, page from ref_text
                # Assuming ref_text format: '(Author, Year, p. Page)'
                ref_text_clean = ref_text.strip('()')
                # Split by comma
                parts = [part.strip() for part in ref_text_clean.split(',')]
                if len(parts) >= 2:
                    author = parts[0]
                    year = parts[1]
                    page = ''
                    if len(parts) > 2:
                        page = parts[2]
                    else:
                        page = ''
                    # Generate citation_key, e.g., 'ridAttributingCyberAttacks2015'
                    # For simplicity, take the first author's last name, first word of title, and year
                    # This is a placeholder; in practice, you'd have a mapping or a more sophisticated method
                    first_author_last_name = author.split()[0].lower()
                    year_number = re.sub(r'\D', '', year)
                    citation_key = f"{first_author_last_name}{year_number}"
                    paragraph_id = ref_id
                    # Reconstruct the reference
                    new_ref = f"([{author}, {year}, {page}]({citation_key}#^{paragraph_id}))"
                    return new_ref
                else:
                    return ref_text
            else:
                return match.group(0)

        # Replace all <ref>...</ref> with the desired format
        paragraph = re.sub(r'<ref.*?>.*?</ref>', replace_ref, paragraph)
        return paragraph

    for entry in data:
        section = entry.get('section', {})

        # Combine introduction, body_text, and conclusion into one HTML content
        html_contents = []
        if 'introduction' in section:
            html_contents.append(section['introduction'])
        if 'body_text' in section:
            html_contents.extend(section['body_text'])
        if 'conclusion' in section:
            html_contents.append(section['conclusion'])

        for html_content in html_contents:
            # Parse the HTML content
            soup = BeautifulSoup(html_content, 'html.parser')

            # Process the content
            # Find all h1 tags
            h1_tags = soup.find_all('h1')
            if not h1_tags:
                continue  # No h1 in this content
            for h1 in h1_tags:
                h1_text = h1.get_text(strip=True)
                # Create a folder named after the h1
                folder_name = h1_text
                os.makedirs(folder_name, exist_ok=True)

                # Collect all elements under this h1 until the next h1
                content_under_h1 = []
                sibling = h1.find_next_sibling()
                while sibling and sibling.name != 'h1':
                    content_under_h1.append(sibling)
                    sibling = sibling.find_next_sibling()

                # Now, process the content under h1
                # We need to process headings and their content
                i = 0
                while i < len(content_under_h1):
                    elem = content_under_h1[i]
                    if elem.name and elem.name.startswith('h') and elem.name != 'h1':
                        heading_level = int(elem.name[1])
                        heading_text = elem.get_text(strip=True)
                        # Create a file named after the heading
                        # Replace invalid filename characters
                        filename_safe = re.sub(r'[\\/*?:"<>|]', "_", heading_text)
                        file_path = os.path.join(folder_name, f"{filename_safe}.md")
                        with open(file_path, 'w', encoding='utf-8') as f:
                            # Write the heading as Markdown
                            heading_md = '#' * heading_level + ' ' + heading_text
                            f.write(heading_md + '\n\n')
                            # Collect content under this heading
                            j = i + 1
                            while j < len(content_under_h1):
                                next_elem = content_under_h1[j]
                                if next_elem.name and next_elem.name.startswith('h'):
                                    next_heading_level = int(next_elem.name[1])
                                    if next_heading_level <= heading_level:
                                        # This is a same or higher level heading, stop collecting
                                        break
                                # Process paragraph or other elements
                                if next_elem.name == 'p':
                                    paragraph_html = str(next_elem)
                                    # Replace <ref> tags
                                    paragraph_html = process_ref_tags(paragraph_html)
                                    # Convert HTML to Markdown
                                    paragraph_md = markdownify.markdownify(paragraph_html, heading_style="ATX")
                                    f.write(paragraph_md + '\n\n')
                                elif next_elem.name and next_elem.name.startswith('h'):
                                    # Skip processing, will be handled in outer loop
                                    pass
                                else:
                                    # Other content
                                    # Convert HTML to Markdown
                                    other_html = str(next_elem)
                                    other_md = markdownify.markdownify(other_html, heading_style="ATX")
                                    f.write(other_md + '\n\n')
                                j += 1
                            i = j - 1  # Update i to continue from the new position
                    i += 1


def export_to_html(data: List[dict], filename: str):
    """
    Export the data to an HTML file by aggregating all HTML content.
    """
    html_content = "<!DOCTYPE html>\n<html>\n<head>\n<meta charset='UTF-8'>\n<title>Exported Section</title>\n</head>\n<body>\n"

    for entry in data:
        section = entry.get('section', {})

        # Introduction
        if 'introduction' in section:
            html_content += section['introduction'] + "\n\n"

        # Body Text
        if 'body_text' in section:
            for body in section['body_text']:
                html_content += body + "\n\n"

        # Conclusion
        if 'conclusion' in section:
            html_content += section['conclusion'] + "\n\n"

    html_content += "</body>\n</html>"

    # Write to HTML file
    with open(f"{filename}.html", "w", encoding='utf-8') as html_file:
        html_file.write(html_content)

def export_to_json(data: List[dict], filename: str):
    """
    Export the data to a JSON file.
    """
    with open(f"{filename}.json", "w", encoding='utf-8') as json_file:
        json.dump(data, json_file, indent=4, ensure_ascii=False)

def export_to_pdf_writen(word_filename: str, pdf_filename: str):
    """
    Convert a Word document to PDF using docx2pdf.
    """
    convert(f"{word_filename}.docx", f"{pdf_filename}.pdf")

def export_written(filename: str, sections: List[dict], format: List[str] = ['md','docx','pdf','html','json']):
    """
    Export the extracted sections data into various formats.

    Parameters:
    - filename: The base name for the exported files.
    - sections: The list of section dictionaries.
    - format: List of formats to export to. Options: 'md', 'docx', 'pdf', 'html', 'json'.
    """
    supported_formats = {'md', 'docx', 'pdf', 'html', 'json'}
    requested_formats = set(format)

    # Validate requested formats
    invalid_formats = requested_formats - supported_formats
    if invalid_formats:
        raise ValueError(f"Unsupported format(s): {', '.join(invalid_formats)}. Supported formats are: {', '.join(supported_formats)}")

    # Export to Markdown
    if 'md' in requested_formats:
        export_to_markdown(sections, filename)
        print(f"Exported to {filename}.md")

    # Export to Word Document
    if 'docx' in requested_formats or 'pdf' in requested_formats:
        export_to_word_written(sections, filename)
        print(f"Exported to {filename}.docx")

    # Export to PDF
    if 'pdf' in requested_formats:
        export_to_pdf_writen(filename, filename)
        print(f"Exported to {filename}.pdf")

    # Export to HTML
    if 'html' in requested_formats:
        export_to_html(sections, filename)
        print(f"Exported to {filename}.html")

    # Export to JSON
    if 'json' in requested_formats:
        export_to_json(sections, filename)
        print(f"Exported to {filename}.json")


def clean_outline(item):
    """Recursively clean the outline to retain only necessary keys, removing 'N/A' and empty values."""
    cleaned_item = {
        'title': item.get('title', ''),
        'level': item.get('level', ''),
        'content': []
    }

    # Only add paragraph-related fields if they are not empty
    paragraph_title = item.get('paragraph_title', [])
    paragraph_text = item.get('paragraph_text', [])
    paragraph_id = item.get('paragraph_id', [])

    if paragraph_title:
        for i in range(len(paragraph_title)):
            cleaned_item['content'].append(
                {'paragraph_title': paragraph_title[i], 'paragraph_quoted': paragraph_text[i],
                 'paragraph_id': paragraph_id[i]})
        cleaned_item['paragraph_count'] = len(paragraph_text)

    # Recursively clean subheadings, removing those without a title
    subheadings = item.get('subheadings', [])
    cleaned_subheadings = [
        clean_outline(subheading) for subheading in subheadings

    ]

    # Only include subheadings if there are any valid ones
    if cleaned_subheadings:
        cleaned_item['subheadings'] = cleaned_subheadings

    return cleaned_item


def clean_themes(themes):
    """Iterate through the themes and clean each outline."""
    cleaned_themes = []
    for theme in themes:
        cleaned_theme = {
            'theme': theme.get('theme', ''),

            'outline': [call_openai_api(data=clean_outline(outline_item),function='writing_sections',id='',model='gpt-4o-2024-08-06',batch=False) for outline_item in theme.get('outline', [])]
        }

        cleaned_themes.append(cleaned_theme)
        with open('theme_writen.txt', 'w', encoding='utf-8') as theme_file:
            theme_file.write(json.dumps(cleaned_themes, indent=4, ensure_ascii=False))
    return cleaned_themes